"""Word .docx 只读抽取纯文本（段落 + 表格）。"""

from __future__ import annotations

import os
from typing import Any

from _meta import skill

_MAX_FILE_BYTES = 40 * 1024 * 1024


def _abs_file(path: str) -> str:
    p = os.path.abspath(os.path.expanduser(path))
    if not os.path.isfile(p):
        raise FileNotFoundError(f"文件不存在: {p}")
    if os.path.getsize(p) > _MAX_FILE_BYTES:
        raise ValueError(f"文件过大（>{_MAX_FILE_BYTES // (1024 * 1024)}MB）。")
    return p


@skill(
    desc="从本地 .docx 抽取段落与表格单元格文本，合并为单一字符串（可设 max_chars）。",
    examples=[
        "docx_to_text('./notes/spec.docx')",
        "docx_to_text('./report.docx', max_chars=50000)",
    ],
)
def docx_to_text(path: str, max_chars: int = 200_000) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少依赖 python-docx，请安装：pip install python-docx") from exc

    abs_path = _abs_file(path)
    if not abs_path.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx")

    max_chars = max(2000, min(int(max_chars), 2_000_000))

    document = Document(abs_path)
    parts: list[str] = []
    for para in document.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                c = (cell.text or "").strip()
                if c:
                    cells.append(c)
            if cells:
                parts.append("\t".join(cells))

    text = "\n".join(parts).strip()
    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars] + "\n…(已截断)"

    return {
        "path": abs_path,
        "char_count": len(text),
        "truncated": truncated,
        "text": text,
    }
